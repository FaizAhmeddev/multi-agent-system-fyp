"""
Finance document export — build PDF, Excel (XLSX), CSV, TXT, and DOCX from
structured finance content. Multiple formats are encoded in parallel (ThreadPoolExecutor).
"""

from __future__ import annotations

import csv
import io
import json
import re
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Callable
from xml.sax.saxutils import escape


def detect_finance_export_intent(text: str) -> bool:
    """True when the user is asking for a downloadable finance/sales document, not only Q&A."""
    if not text or len(text.strip()) < 8:
        return False
    low = text.lower()
    # Output format mentioned (word boundary avoids matching unrelated substrings)
    has_format = bool(
        re.search(
            r"\b(pdf|excel|xlsx|xls|csv|docx|spreadsheet|word)\b|\.pdf\b|\.xlsx\b|\.csv\b|\.docx\b",
            low,
        )
    )
    has_action = any(
        a in low
        for a in (
            "generate",
            "export",
            "create",
            "download",
            "save as",
            "build",
            "make ",
            "produce",
            "output",
        )
    )
    # Markdown / pipe table often pasted with "given data" prompts
    has_table = "|" in text and "\n" in text and ("---" in text or re.search(r"\|\s*[^|\n]+\s*\|", text))
    fin_sales_kw = (
        "finance",
        "financial",
        "budget",
        "expense",
        "invoice",
        "revenue",
        "cash flow",
        "p&l",
        "profit",
        "loss",
        "pkr",
        "usd",
        "report",
        "summary",
        "forecast",
        "ledger",
        "payroll",
        "statement",
        "sales",
        "sold",
        "quantity",
        "inventory",
        "transaction",
        "given data",
        "as per",
    )
    has_domain = any(k in low for k in fin_sales_kw)

    export_kw = (
        "generate pdf",
        "generate excel",
        "generate xlsx",
        "generate xlxs",
        "generate xls",
        "generate csv",
        "generate docx",
        "generate word",
        "generate txt",
        "generate summary",
        "generate report",
        "export pdf",
        "export excel",
        "export xlsx",
        "export csv",
        "download pdf",
        "download excel",
        "save as pdf",
        "save as excel",
        "create pdf",
        "create excel",
        "build pdf",
        "build excel",
        "write pdf",
        "write excel",
        "make pdf",
        "make excel",
        "spreadsheet",
        " in pdf",
        " in excel",
        " in xlsx",
        " in csv",
        " in docx",
        "summary in pdf",
        "report in pdf",
        ".pdf",
        ".xlsx",
        ".csv",
        ".docx",
    )
    strict_export = any(k in low for k in export_kw)

    # Original rule: explicit export phrase + finance language
    if strict_export and has_domain:
        return True
    # "Generate summary in PDF" + pasted sales table (no word "finance")
    if has_format and has_action and (has_domain or has_table):
        return True
    # Table + clear PDF/Excel target (e.g. "… data … summary … pdf")
    if has_table and has_format and ("summary" in low or "report" in low or has_action):
        return True
    return False


def infer_formats_from_text(text: str, explicit: list[str] | None = None) -> list[str]:
    """Return canonical format slugs: pdf, xlsx, csv, txt, docx."""
    canon = {"pdf", "xlsx", "csv", "txt", "docx"}
    if explicit:
        alias = {
            "excel": "xlsx",
            "word": "docx",
            "xlxs": "xlsx",
            "xlx": "xlsx",
            "xls": "xlsx",
            ".xls": "xlsx",
        }
        out: list[str] = []
        for f in explicit:
            raw = f.strip().lower()
            x = alias.get(raw, raw)
            if x in canon:
                out.append(x)
        return list(dict.fromkeys(out)) or ["pdf"]

    low = (text or "").lower()
    found: list[str] = []
    if "pdf" in low or ".pdf" in low:
        found.append("pdf")
    if "xlsx" in low or "excel" in low or "spreadsheet" in low or ".xlsx" in low or " xls" in low:
        found.append("xlsx")
    if "csv" in low or ".csv" in low:
        found.append("csv")
    if "docx" in low or "word document" in low or ".docx" in low:
        found.append("docx")
    if "txt" in low or "text file" in low or ".txt" in low:
        found.append("txt")
    if not found:
        found = ["pdf"]
    return list(dict.fromkeys([f for f in found if f in canon]))


def _strip_json_fence(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.I)
    s = re.sub(r"\s*```\s*$", "", s)
    return s.strip()


def llm_finance_document_structure(
    *,
    user_request: str,
    source_data: str,
    user_name: str,
) -> dict[str, Any]:
    """Ask the LLM for a JSON structure used to populate all export formats."""
    import os
    from langchain_openai import ChatOpenAI
    from config import OPENAI_API_KEY

    os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.15)
    prompt = f"""You are a Finance Document Generator. The user wants exportable finance documents.

User: {user_name}
Request:
{user_request[:3500]}

Source data / context (may be empty):
{source_data[:6000]}

Return ONLY valid JSON (no markdown fences) with this exact shape:
{{
  "title": "string, document title",
  "document_type": "summary|report|invoice_analysis|expense_analysis|budget|custom",
  "executive_summary": "string, 2-4 sentences",
  "sections": [{{"heading": "string", "body": "string (plain text, can use newlines)"}}],
  "tables": [{{"name": "string", "headers": ["col1","col2"], "rows": [["a","b"],["c","d"]]}}],
  "key_metrics": [{{"label": "string", "value": "string"}}],
  "recommendations": ["string", "string"]
}}

Rules:
- Use at least 2 sections and 1 table when source_data suggests numbers or line items; otherwise still include 1 small table (e.g. Key metrics).
- Amounts: prefer PKR or USD labels in cells when relevant.
- Keep each section body under 1200 characters; tables max 30 rows.
"""
    raw = llm.invoke(prompt).content or "{}"
    try:
        return json.loads(_strip_json_fence(raw))
    except json.JSONDecodeError:
        return {
            "title": "Finance document",
            "document_type": "custom",
            "executive_summary": raw[:800],
            "sections": [{"heading": "Details", "body": raw[:4000]}],
            "tables": [{"name": "Notes", "headers": ["Item", "Value"], "rows": [["Summary", raw[:500]]]}],
            "key_metrics": [],
            "recommendations": [],
        }


def _flatten_body(data: dict[str, Any]) -> str:
    lines: list[str] = [data.get("title") or "Finance document", "", data.get("executive_summary") or "", ""]
    for s in data.get("sections") or []:
        if isinstance(s, dict):
            lines.append(str(s.get("heading", "")))
            lines.append(str(s.get("body", "")))
            lines.append("")
    for rec in data.get("recommendations") or []:
        lines.append(f"- {rec}")
    return "\n".join(lines).strip()


def _build_pdf(data: dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib import colors

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=data.get("title") or "Finance")
    styles = getSampleStyleSheet()
    story: list[Any] = []
    title = escape(str(data.get("title") or "Finance document"))
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 10))
    es = str(data.get("executive_summary") or "")
    if es:
        story.append(Paragraph(escape(es).replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 12))
    for s in data.get("sections") or []:
        if not isinstance(s, dict):
            continue
        h = escape(str(s.get("heading", "")))
        b = escape(str(s.get("body", ""))).replace("\n", "<br/>")
        story.append(Paragraph(f"<b>{h}</b>", styles["Heading2"]))
        story.append(Paragraph(b, styles["Normal"]))
        story.append(Spacer(1, 10))
    for tbl in data.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        headers = [str(x) for x in (tbl.get("headers") or ["Column A", "Column B"])]
        rows_raw = tbl.get("rows") or []
        tdata = [headers]
        for row in rows_raw[:40]:
            if isinstance(row, (list, tuple)):
                tdata.append([str(x) for x in row])
            else:
                tdata.append([str(row)])
        if len(tdata) > 1:
            tw = Table(tdata, repeatRows=1)
            tw.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#166534")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0fdf4")]),
                    ]
                )
            )
            story.append(Spacer(1, 8))
            story.append(Paragraph(f"<b>{escape(str(tbl.get('name', 'Table')))}</b>", styles["Heading3"]))
            story.append(tw)
            story.append(Spacer(1, 12))
    doc.build(story)
    return buf.getvalue()


def _build_xlsx(data: dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Overview"
    ws["A1"] = data.get("title") or "Finance document"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A3"] = "Executive summary"
    ws["A3"].font = Font(bold=True)
    ws["A4"] = data.get("executive_summary") or ""
    ws["A4"].alignment = Alignment(wrap_text=True)
    r = 6
    for s in data.get("sections") or []:
        if not isinstance(s, dict):
            continue
        ws.cell(row=r, column=1, value=str(s.get("heading", ""))).font = Font(bold=True)
        r += 1
        ws.cell(row=r, column=1, value=str(s.get("body", ""))).alignment = Alignment(wrap_text=True)
        r += 2
    for km in data.get("key_metrics") or []:
        if isinstance(km, dict):
            ws.cell(row=r, column=1, value=str(km.get("label", "")))
            ws.cell(row=r, column=2, value=str(km.get("value", "")))
            r += 1
    r += 1
    ws.cell(row=r, column=1, value="Recommendations").font = Font(bold=True)
    r += 1
    for rec in data.get("recommendations") or []:
        ws.cell(row=r, column=1, value=str(rec))
        r += 1

    for tbl in data.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        name = re.sub(r"[^\w\- ]", "", str(tbl.get("name") or "Data"))[:28] or "Data"
        sh = wb.create_sheet(title=name[:31])
        headers = [str(x) for x in (tbl.get("headers") or ["A", "B"])]
        for c, h in enumerate(headers, 1):
            cell = sh.cell(row=1, column=c, value=h)
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="DCFCE7")
        for ri, row in enumerate((tbl.get("rows") or [])[:200], 2):
            if isinstance(row, (list, tuple)):
                for ci, val in enumerate(row, 1):
                    sh.cell(row=ri, column=ci, value=val)
            else:
                sh.cell(row=ri, column=1, value=str(row))

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def _build_csv(data: dict[str, Any]) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Finance export"])
    w.writerow([data.get("title") or ""])
    w.writerow([])
    w.writerow(["Executive summary"])
    w.writerow([data.get("executive_summary") or ""])
    w.writerow([])
    for s in data.get("sections") or []:
        if isinstance(s, dict):
            w.writerow([s.get("heading", "")])
            w.writerow([s.get("body", "")])
            w.writerow([])
    for tbl in data.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        w.writerow([str(tbl.get("name", "Table"))])
        headers = [str(x) for x in (tbl.get("headers") or [])]
        if headers:
            w.writerow(headers)
        for row in (tbl.get("rows") or [])[:200]:
            if isinstance(row, (list, tuple)):
                w.writerow([str(x) for x in row])
            else:
                w.writerow([str(row)])
        w.writerow([])
    return buf.getvalue().encode("utf-8-sig")


def _build_txt(data: dict[str, Any]) -> bytes:
    return _flatten_body(data).encode("utf-8")


def _build_docx(data: dict[str, Any]) -> bytes:
    from docx import Document

    d = Document()
    d.add_heading(str(data.get("title") or "Finance document"), 0)
    if data.get("executive_summary"):
        d.add_paragraph(str(data.get("executive_summary")))
    for s in data.get("sections") or []:
        if isinstance(s, dict):
            d.add_heading(str(s.get("heading", "")), level=2)
            d.add_paragraph(str(s.get("body", "")))
    for tbl in data.get("tables") or []:
        if not isinstance(tbl, dict):
            continue
        d.add_heading(str(tbl.get("name", "Table")), level=3)
        headers = [str(x) for x in (tbl.get("headers") or ["A", "B"])]
        table = d.add_table(rows=1, cols=len(headers))
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
        for row in (tbl.get("rows") or [])[:50]:
            cells = table.add_row().cells
            if isinstance(row, (list, tuple)):
                vals = [str(x) for x in row]
                while len(vals) < len(headers):
                    vals.append("")
                for i in range(len(headers)):
                    cells[i].text = vals[i] if i < len(vals) else ""
            else:
                cells[0].text = str(row)
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


_BUILDERS: dict[str, Callable[[dict[str, Any]], bytes]] = {
    "pdf": _build_pdf,
    "xlsx": _build_xlsx,
    "csv": _build_csv,
    "txt": _build_txt,
    "docx": _build_docx,
}


def export_finance_documents_parallel(structured: dict[str, Any], formats: list[str]) -> list[dict[str, Any]]:
    """Build each requested format in parallel; returns list of dicts with filename, data (bytes), mime_type."""
    formats = [f for f in formats if f in _BUILDERS]
    if not formats:
        formats = ["pdf"]
    slug = uuid.uuid4().hex[:8]
    base = re.sub(r"[^\w\- ]", "", str(structured.get("title") or "finance_export"))[:40].strip() or "finance_export"
    base = base.replace(" ", "_")

    def _one(fmt: str) -> dict[str, Any]:
        builder = _BUILDERS[fmt]
        data = builder(structured)
        ext = fmt
        mime = {
            "pdf": "application/pdf",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "csv": "text/csv",
            "txt": "text/plain",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }[fmt]
        return {
            "filename": f"{base}_{slug}.{ext}",
            "data": data,
            "mime_type": mime,
            "format": fmt,
        }

    out: list[dict[str, Any]] = []
    with ThreadPoolExecutor(max_workers=min(5, len(formats))) as ex:
        futs = {ex.submit(_one, fmt): fmt for fmt in formats}
        for fut in as_completed(futs):
            out.append(fut.result())
    # stable order: pdf, xlsx, csv, txt, docx
    order = {f: i for i, f in enumerate(["pdf", "xlsx", "csv", "txt", "docx"])}
    out.sort(key=lambda x: order.get(x.get("format", ""), 99))
    return out


def run_finance_document_export(
    *,
    user_request: str,
    source_data: str,
    user_name: str,
    export_formats: list[str] | None = None,
) -> dict[str, Any]:
    """
    Full pipeline: structured content via LLM → parallel file builds.
    Returns: ok, output (markdown summary for UI), export_files (list of file dicts), error (optional).
    """
    try:
        formats = infer_formats_from_text(user_request, export_formats)
        structured = llm_finance_document_structure(
            user_request=user_request,
            source_data=source_data or "",
            user_name=user_name or "User",
        )
        files = export_finance_documents_parallel(structured, formats)
        lines = [
            "### Finance documents generated",
            "",
            f"**Title:** {structured.get('title', '')}",
            f"**Type:** {structured.get('document_type', '')}",
            f"**Formats (parallel):** {', '.join(formats)}",
            "",
            structured.get("executive_summary", ""),
            "",
            "Download the files below.",
        ]
        return {
            "ok": True,
            "output": "\n".join(lines),
            "export_files": files,
            "structured": structured,
        }
    except Exception as e:
        return {
            "ok": False,
            "output": f"❌ **Document export failed:** {e}",
            "export_files": [],
            "error": str(e),
        }


def finance_export_summary_markdown(structured: dict[str, Any], formats: list[str]) -> str:
    """Short markdown for chat when files are attached separately."""
    parts = [
        f"**{structured.get('title', 'Finance document')}** ({structured.get('document_type', '')})",
        f"**Files:** {', '.join(formats)}",
        "",
        str(structured.get("executive_summary", "")),
    ]
    return "\n".join(parts)
